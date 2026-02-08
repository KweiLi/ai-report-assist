"""Step 6: Export debiased report to PDF or DOCX."""

from __future__ import annotations

import logging
from collections import Counter
from datetime import datetime
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF

from app.pipeline.unmasker import unmask_phrase

logger = logging.getLogger(__name__)

# Bias type display colours (hex without #) – mirrors BIAS_COLORS in debiaser.py
_BIAS_HEX = {
    "RACIAL_ETHNIC": "e74c3c",
    "GENDER": "e67e22",
    "SOCIOECONOMIC": "9b59b6",
    "STEREOTYPING": "f39c12",
    "INFLAMMATORY": "c0392b",
    "CONFIRMATION": "3498db",
    "SUBJECTIVE": "e84393",
}

_BIAS_LABELS = {
    "RACIAL_ETHNIC": "Racial / Ethnic",
    "GENDER": "Gender",
    "SOCIOECONOMIC": "Socioeconomic",
    "STEREOTYPING": "Stereotyping",
    "INFLAMMATORY": "Inflammatory",
    "CONFIRMATION": "Confirmation",
    "SUBJECTIVE": "Subjective",
}


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _lighten_hex(hex_color: str, factor: float = 0.85) -> str:
    """Lighten a hex colour for use as a background tint."""
    r, g, b = _hex_to_rgb(hex_color)
    r = int(r + (255 - r) * factor)
    g = int(g + (255 - g) * factor)
    b = int(b + (255 - b) * factor)
    return f"{r:02x}{g:02x}{b:02x}"


# ── DOCX helpers ──

def _set_cell_shading(cell, hex_color: str) -> None:
    """Set background shading colour on a DOCX table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.upper())
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_border_left(cell, hex_color: str, width: int = 12) -> None:
    """Set a coloured left border on a DOCX table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width))
    left.set(qn("w:color"), hex_color.upper())
    left.set(qn("w:space"), "0")
    existing = borders.find(qn("w:left"))
    if existing is not None:
        borders.remove(existing)
    borders.append(left)


def _hide_cell_borders(cell) -> None:
    """Remove all borders except left from a DOCX table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("w:top", "w:right", "w:bottom"):
        el = OxmlElement(side)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "FFFFFF")
        existing = borders.find(qn(side))
        if existing is not None:
            borders.remove(existing)
        borders.append(el)


def _docx_bias_summary_table(doc: Document, bias_changes: list[dict]) -> None:
    """Add a colour-coded summary table of bias type counts."""
    counts = Counter(c["bias_type"] for c in bias_changes)
    table = doc.add_table(rows=1, cols=len(counts))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (bias_type, count) in enumerate(counts.most_common()):
        cell = table.cell(0, i)
        color = _BIAS_HEX.get(bias_type, "999999")
        _set_cell_shading(cell, _lighten_hex(color, 0.75))
        _set_cell_border_left(cell, color, width=18)
        p = cell.paragraphs[0]
        label_run = p.add_run(_BIAS_LABELS.get(bias_type, bias_type))
        label_run.font.size = Pt(9)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(*_hex_to_rgb(color))
        count_run = p.add_run(f"  ({count})")
        count_run.font.size = Pt(8)
        count_run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()  # spacer


def _docx_bias_change_card(doc: Document, change: dict) -> None:
    """Add a single colour-coded bias change card as a 1-cell table."""
    color = _BIAS_HEX.get(change["bias_type"], "999999")

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    _set_cell_border_left(cell, color, width=18)
    _hide_cell_borders(cell)
    _set_cell_shading(cell, "FAFAFA")

    # Bias type tag
    p = cell.paragraphs[0]
    tag_run = p.add_run(f"  {_BIAS_LABELS.get(change['bias_type'], change['bias_type']).upper()}  ")
    tag_run.font.size = Pt(7)
    tag_run.font.bold = True
    tag_run.font.color.rgb = RGBColor(255, 255, 255)
    # Simulate tag background via shading on the run
    rpr = tag_run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color.upper())
    rpr.append(shd)

    # Original phrase
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(1)
    orig_label = p2.add_run("Original: ")
    orig_label.font.size = Pt(9)
    orig_label.font.bold = True
    orig_text = p2.add_run(f'"{change["original_phrase"]}"')
    orig_text.font.size = Pt(9)
    orig_text.font.color.rgb = RGBColor(192, 57, 43)  # red
    orig_text.font.strike = True

    # Replacement phrase
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(1)
    p3.paragraph_format.space_after = Pt(1)
    repl_label = p3.add_run("Replacement: ")
    repl_label.font.size = Pt(9)
    repl_label.font.bold = True
    repl_text = p3.add_run(f'"{change["replacement_phrase"]}"')
    repl_text.font.size = Pt(9)
    repl_text.font.color.rgb = RGBColor(39, 174, 96)  # green
    repl_text.font.bold = True

    # Explanation
    p4 = cell.add_paragraph()
    p4.paragraph_format.space_before = Pt(2)
    p4.paragraph_format.space_after = Pt(2)
    expl_run = p4.add_run(change["explanation"])
    expl_run.font.size = Pt(8)
    expl_run.font.italic = True
    expl_run.font.color.rgb = RGBColor(100, 100, 100)


# ── Main export functions ──

def export_docx(
    text: str,
    output_path: str | Path,
    title: str = "Debiased Report",
    entities_found: list[dict] | None = None,
    changes_summary: str | None = None,
    bias_changes: list[dict] | None = None,
    acronyms_preserved: list[dict] | None = None,
) -> Path:
    """Export debiased report with coloured bias analysis to Word."""
    output_path = Path(output_path)
    doc = Document()

    # Title
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.size = Pt(16)

    # Metadata
    meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    meta.runs[0].font.size = Pt(9)

    # --- Section 1: Debiased Report ---
    doc.add_heading("Debiased Report", level=2)
    for paragraph in text.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)

    # --- Section 2: Bias Analysis (colour-coded) ---
    if bias_changes:
        doc.add_page_break()
        doc.add_heading(f"Bias Analysis - {len(bias_changes)} Issues Found", level=2)

        intro = doc.add_paragraph(
            "The following biased phrases were identified and corrected. "
            "Each item shows the bias type, original phrasing, neutral replacement, "
            "and an explanation."
        )
        intro.runs[0].font.italic = True
        intro.runs[0].font.size = Pt(9)
        intro.runs[0].font.color.rgb = RGBColor(100, 100, 100)

        # Summary count bar
        _docx_bias_summary_table(doc, bias_changes)

        # Individual change cards
        for change in bias_changes:
            _docx_bias_change_card(doc, change)

    # --- Section 3: Masking Details ---
    if entities_found:
        doc.add_page_break()
        doc.add_heading("Masking Details", level=2)
        intro = doc.add_paragraph(
            f"{len(entities_found)} sensitive entities were detected and masked "
            "before sending the report to the AI for debiasing. "
            "Only masked text was transmitted to the cloud."
        )
        intro.runs[0].font.italic = True
        intro.runs[0].font.size = Pt(9)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, label in enumerate(["Type", "Original Value", "Masked Token", "Confidence"]):
            hdr[i].text = label
            hdr[i].paragraphs[0].runs[0].font.bold = True
            hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

        for entity in entities_found:
            row = table.add_row().cells
            row[0].text = entity["entity_type"]
            row[1].text = entity["original"]
            row[2].text = entity["token"]
            row[3].text = str(entity["score"])
            for cell in row:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

    # --- Section 4: Acronyms & Abbreviations ---
    if acronyms_preserved:
        doc.add_page_break()
        doc.add_heading("Acronyms & Abbreviations", level=2)
        intro = doc.add_paragraph(
            f"{len(acronyms_preserved)} acronyms/abbreviations were identified in the report. "
            "These are standard law enforcement terminology and were intentionally "
            "preserved (not masked) during processing."
        )
        intro.runs[0].font.italic = True
        intro.runs[0].font.size = Pt(9)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, label in enumerate(["Acronym", "Initially Detected As", "Action"]):
            hdr[i].text = label
            hdr[i].paragraphs[0].runs[0].font.bold = True
            hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

        for acr in acronyms_preserved:
            row = table.add_row().cells
            row[0].text = acr["text"]
            row[1].text = acr["detected_as"]
            row[2].text = acr["reason"]
            for cell in row:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.save(str(output_path))
    logger.info("Exported DOCX to %s", output_path)
    return output_path


# ── PDF helpers ──

def _pdf_section_heading(pdf: FPDF, text: str) -> None:
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(30, 30, 60)
    pdf.cell(0, 9, text, new_x="LMARGIN", new_y="NEXT")
    pdf.set_draw_color(30, 30, 60)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(4)


def _pdf_bias_summary_bar(pdf: FPDF, bias_changes: list[dict]) -> None:
    """Draw a row of colour-coded bias type count chips."""
    counts = Counter(c["bias_type"] for c in bias_changes)
    x_start = pdf.l_margin
    for bias_type, count in counts.most_common():
        color = _BIAS_HEX.get(bias_type, "999999")
        r, g, b = _hex_to_rgb(color)
        lr, lg, lb = _hex_to_rgb(_lighten_hex(color, 0.75))
        label = _BIAS_LABELS.get(bias_type, bias_type)
        chip_text = f" {label} ({count}) "

        pdf.set_font("Helvetica", "B", 8)
        chip_w = pdf.get_string_width(chip_text) + 4

        # Check if chip fits on current line
        if pdf.get_x() + chip_w > pdf.w - pdf.r_margin:
            pdf.ln(8)

        # Draw chip background
        pdf.set_fill_color(lr, lg, lb)
        pdf.set_draw_color(r, g, b)
        chip_x = pdf.get_x()
        chip_y = pdf.get_y()
        pdf.rect(chip_x, chip_y, chip_w, 7, style="FD")
        # Coloured left accent
        pdf.set_fill_color(r, g, b)
        pdf.rect(chip_x, chip_y, 2, 7, style="F")

        pdf.set_text_color(r, g, b)
        pdf.set_xy(chip_x, chip_y)
        pdf.cell(chip_w, 7, chip_text)
        pdf.set_x(pdf.get_x() + 3)  # gap between chips

    pdf.ln(10)


def _pdf_bias_change_card(pdf: FPDF, change: dict) -> None:
    """Draw a single colour-coded bias change card in the PDF."""
    color = _BIAS_HEX.get(change["bias_type"], "999999")
    r, g, b = _hex_to_rgb(color)

    card_x = pdf.l_margin
    card_w = pdf.w - pdf.l_margin - pdf.r_margin
    card_start_y = pdf.get_y()

    # Check if we need a new page (estimate ~28mm per card)
    if card_start_y + 28 > pdf.h - pdf.b_margin:
        pdf.add_page()
        card_start_y = pdf.get_y()

    # Card background
    pdf.set_fill_color(250, 250, 250)
    pdf.set_draw_color(230, 230, 230)

    # We'll draw the background after calculating height, so save position
    content_x = card_x + 5  # indent past left accent bar

    # Bias type tag
    pdf.set_xy(content_x, card_start_y + 2)
    tag_label = f"  {_BIAS_LABELS.get(change['bias_type'], change['bias_type']).upper()}  "
    pdf.set_font("Helvetica", "B", 7)
    tag_w = pdf.get_string_width(tag_label) + 2
    pdf.set_fill_color(r, g, b)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(tag_w, 5, tag_label, fill=True)
    pdf.ln(6)

    # Original phrase
    pdf.set_x(content_x)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(18, 5, "Original: ")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(192, 57, 43)
    orig_text = _sanitize_for_pdf(f'"{change["original_phrase"]}"')
    pdf.multi_cell(card_w - 23, 5, orig_text)

    # Replacement phrase
    pdf.set_x(content_x)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(24, 5, "Replacement: ")
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(39, 174, 96)
    repl_text = _sanitize_for_pdf(f'"{change["replacement_phrase"]}"')
    pdf.multi_cell(card_w - 29, 5, repl_text)

    # Explanation
    pdf.set_x(content_x)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(100, 100, 100)
    explanation = _sanitize_for_pdf(change["explanation"])
    pdf.multi_cell(card_w - 10, 4, explanation)

    card_end_y = pdf.get_y() + 2

    # Draw card border and accent bar over the content
    pdf.set_fill_color(250, 250, 250)
    pdf.set_draw_color(230, 230, 230)
    pdf.rect(card_x, card_start_y, card_w, card_end_y - card_start_y, style="D")
    # Left accent bar
    pdf.set_fill_color(r, g, b)
    pdf.rect(card_x, card_start_y, 3, card_end_y - card_start_y, style="F")

    pdf.set_y(card_end_y + 3)


def export_pdf(
    text: str,
    output_path: str | Path,
    title: str = "Debiased Report",
    entities_found: list[dict] | None = None,
    changes_summary: str | None = None,
    bias_changes: list[dict] | None = None,
    acronyms_preserved: list[dict] | None = None,
) -> Path:
    """Export debiased report with coloured bias analysis to PDF."""
    output_path = Path(output_path)
    title = _sanitize_for_pdf(title)
    text = _sanitize_for_pdf(text)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    # --- Page 1: Debiased Report ---
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 6, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    _pdf_section_heading(pdf, "Debiased Report")
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 10)
    for paragraph in text.split("\n"):
        if paragraph.strip():
            pdf.multi_cell(0, 5, paragraph)
            pdf.ln(2)

    # --- Page 2: Bias Analysis (colour-coded) ---
    if bias_changes:
        pdf.add_page()
        _pdf_section_heading(pdf, f"Bias Analysis - {len(bias_changes)} Issues Found")

        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(
            0, 5,
            "The following biased phrases were identified and corrected. "
            "Each item shows the bias type, original phrasing, neutral replacement, "
            "and an explanation.",
        )
        pdf.ln(4)

        # Summary bar
        _pdf_bias_summary_bar(pdf, bias_changes)

        # Individual change cards
        for change in bias_changes:
            _pdf_bias_change_card(pdf, change)

    # --- Page 3: Masking Details ---
    if entities_found:
        pdf.add_page()
        _pdf_section_heading(pdf, "Masking Details")
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(
            0, 5,
            f"{len(entities_found)} sensitive entities were detected and masked "
            "before sending the report to the AI for debiasing. "
            "Only masked text was transmitted to the cloud.",
        )
        pdf.ln(4)
        pdf.set_text_color(0, 0, 0)

        col_widths = [45, 55, 55, 25]
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(230, 230, 230)
        for w, label in zip(col_widths, ["Type", "Original Value", "Masked Token", "Confidence"]):
            pdf.cell(w, 7, label, border=1, fill=True)
        pdf.ln()

        pdf.set_font("Helvetica", "", 8)
        for entity in entities_found:
            row_data = [
                entity["entity_type"],
                _truncate(_sanitize_for_pdf(entity["original"]), 30),
                entity["token"],
                str(entity["score"]),
            ]
            for w, val in zip(col_widths, row_data):
                pdf.cell(w, 6, val, border=1)
            pdf.ln()

    # --- Page 4: Acronyms & Abbreviations ---
    if acronyms_preserved:
        pdf.add_page()
        _pdf_section_heading(pdf, "Acronyms & Abbreviations")
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(
            0, 5,
            f"{len(acronyms_preserved)} acronyms/abbreviations were identified in the report. "
            "These are standard law enforcement terminology and were intentionally "
            "preserved (not masked) during processing.",
        )
        pdf.ln(4)
        pdf.set_text_color(0, 0, 0)

        col_widths = [30, 50, 100]
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(230, 230, 230)
        for w, label in zip(col_widths, ["Acronym", "Detected As", "Action"]):
            pdf.cell(w, 7, label, border=1, fill=True)
        pdf.ln()

        pdf.set_font("Helvetica", "", 8)
        for acr in acronyms_preserved:
            pdf.cell(30, 6, _sanitize_for_pdf(acr["text"]), border=1)
            pdf.cell(50, 6, _sanitize_for_pdf(acr["detected_as"]), border=1)
            pdf.cell(100, 6, _truncate(_sanitize_for_pdf(acr["reason"]), 55), border=1)
            pdf.ln()

    pdf.output(str(output_path))
    logger.info("Exported PDF to %s", output_path)
    return output_path


def _sanitize_for_pdf(text: str) -> str:
    """Replace Unicode characters that fpdf2's built-in fonts can't encode."""
    replacements = {
        "\u2018": "'",   # left single quote
        "\u2019": "'",   # right single quote
        "\u201c": '"',   # left double quote
        "\u201d": '"',   # right double quote
        "\u2013": "-",   # en dash
        "\u2014": "-",   # em dash
        "\u2026": "...", # ellipsis
        "\u00a0": " ",   # non-breaking space
        "\u2022": "*",   # bullet
        "\u200b": "",    # zero-width space
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text


def _truncate(text: str, max_len: int) -> str:
    return text if len(text) <= max_len else text[: max_len - 3] + "..."


# ── Formatted PDF export (preserve original layout) ──


def _ocr_page_blocks(page: fitz.Page, dpi: int = 300) -> list[dict]:
    """OCR a single page and return text blocks with bounding rects.

    Each block dict has: block_num, text, rect (fitz.Rect in PDF points),
    word_rects (list of per-word fitz.Rect), avg_word_height (PDF points).
    """
    import pytesseract
    from PIL import Image
    from pytesseract import Output

    scale = 72.0 / dpi
    pix = page.get_pixmap(dpi=dpi)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    ocr_data = pytesseract.image_to_data(img, output_type=Output.DICT)

    # Group word indices by block_num
    block_indices: dict[int, list[int]] = {}
    for i in range(len(ocr_data["text"])):
        if not ocr_data["text"][i].strip():
            continue
        block_indices.setdefault(ocr_data["block_num"][i], []).append(i)

    blocks: list[dict] = []
    for bnum, indices in block_indices.items():
        word_rects = []
        for i in indices:
            word_rects.append(fitz.Rect(
                ocr_data["left"][i] * scale,
                ocr_data["top"][i] * scale,
                (ocr_data["left"][i] + ocr_data["width"][i]) * scale,
                (ocr_data["top"][i] + ocr_data["height"][i]) * scale,
            ))
        block_rect = fitz.Rect(
            min(r.x0 for r in word_rects),
            min(r.y0 for r in word_rects),
            max(r.x1 for r in word_rects),
            max(r.y1 for r in word_rects),
        )
        avg_h = sum(r.height for r in word_rects) / len(word_rects)
        blocks.append({
            "block_num": bnum,
            "text": " ".join(ocr_data["text"][i] for i in indices),
            "rect": block_rect,
            "word_rects": word_rects,
            "avg_word_height": avg_h,
        })

    # Sort top-to-bottom, then left-to-right
    blocks.sort(key=lambda b: (b["rect"].y0, b["rect"].x0))
    return blocks


def _normalize_ws(text: str) -> str:
    """Collapse all whitespace (spaces, newlines, tabs) to single spaces."""
    return " ".join(text.split())


def _merge_blocks(blocks: list[dict]) -> dict:
    """Merge multiple OCR blocks into a single virtual block."""
    all_word_rects: list[fitz.Rect] = []
    for b in blocks:
        all_word_rects.extend(b["word_rects"])
    return {
        "block_num": blocks[0]["block_num"],
        "text": " ".join(b["text"] for b in blocks),
        "rect": fitz.Rect(
            min(r.x0 for r in all_word_rects),
            min(r.y0 for r in all_word_rects),
            max(r.x1 for r in all_word_rects),
            max(r.y1 for r in all_word_rects),
        ),
        "word_rects": all_word_rects,
        "avg_word_height": sum(r.height for r in all_word_rects) / len(all_word_rects),
    }


def _estimate_textbox_height(num_lines: int, fontsize: float) -> float:
    """Estimate the height needed by insert_textbox for N lines of text.

    Measured from PyMuPDF's actual rendering behaviour:
      first line height  ≈ fontsize × 1.7
      subsequent spacing ≈ fontsize × 1.4
    """
    if num_lines <= 0:
        return 0
    return fontsize * 1.7 + max(0, num_lines - 1) * fontsize * 1.4


def _whiteout_and_render(
    page: fitz.Page,
    block: dict,
    text: str,
    fontname: str = "helv",
) -> None:
    """White out a block area and render replacement text.

    If the replacement text is longer than the original, the render
    rectangle is **extended downward** so all text is visible at a
    readable font size.  The extended area is also whited-out.
    """
    rect = block["rect"]
    margin = 2

    # Font size: derived from OCR measurement, never below 7 pt
    fontsize = max(block["avg_word_height"] * 0.80, 7)

    # If the text is much longer, allow a *small* font reduction (down to 7 pt)
    # but primarily rely on extending the rect.
    #
    # PyMuPDF insert_textbox actual line heights (measured):
    #   first line  ≈ fontsize × 1.7
    #   subsequent  ≈ fontsize × 1.4
    # Using these values (not the old 1.3) to correctly size the render rect.
    wrapped = _wrap_text(text.strip(), rect.width, fontname, fontsize)
    needed = _estimate_textbox_height(len(wrapped), fontsize)

    if needed > rect.height * 1.1 and fontsize > 7:
        # Try one modest reduction (max 20 %) before extending
        test_fs = max(fontsize * 0.80, 7)
        test_wrap = _wrap_text(text.strip(), rect.width, fontname, test_fs)
        test_needed = _estimate_textbox_height(len(test_wrap), test_fs)
        if test_needed <= rect.height * 1.1:
            fontsize = test_fs
            wrapped = test_wrap
            needed = test_needed

    # Determine render rect — extend downward if necessary
    # Add a small buffer (4pt) to prevent bottom-line clipping
    render_rect = fitz.Rect(
        rect.x0, rect.y0, rect.x1, max(rect.y1, rect.y0 + needed + 4),
    )

    # White out the entire area (original block + any extension)
    clean = fitz.Rect(
        render_rect.x0 - margin, render_rect.y0 - margin,
        render_rect.x1 + margin, render_rect.y1 + margin,
    )
    page.draw_rect(clean, color=None, fill=(1, 1, 1))

    # Render the debiased text
    rc = page.insert_textbox(
        render_rect, text.strip(),
        fontsize=fontsize, fontname=fontname, color=(0, 0, 0),
    )

    logger.info(
        "Rendered %d chars in %.0fx%.0f rect (fontsize=%.1f, extended=%s, overflow=%s)",
        len(text), render_rect.width, render_rect.height,
        fontsize, render_rect.height > rect.height + 1, rc < 0,
    )


def _wrap_text(
    text: str, max_width: float, fontname: str, fontsize: float
) -> list[str]:
    """Word-wrap text into lines that fit within max_width."""
    lines: list[str] = []
    for paragraph in text.split("\n"):
        if not paragraph.strip():
            lines.append("")
            continue
        words = paragraph.split()
        current = ""
        for word in words:
            test = current + (" " if current else "") + word
            w = fitz.get_text_length(test, fontname=fontname, fontsize=fontsize)
            if w <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines


def _replace_affected_blocks(
    doc: fitz.Document,
    changes: list[tuple[str, str]],
    page_texts: list[str] | None = None,
) -> None:
    """Block-level replacement using **stored text** for content and OCR
    only for spatial positions.

    Approach — paragraph-first:
    1. Split the stored page text into paragraphs.
    2. For each paragraph, apply bias changes.  Skip unchanged paragraphs.
    3. For each changed paragraph, find the OCR blocks that correspond to
       it using word-set overlap (robust to OCR character errors).
    4. Merge matching blocks, white-out, render the debiased **stored**
       paragraph.  No OCR text is ever rendered, so character errors like
       S→$ or I→| are avoided.
    """
    import re

    for page_idx, page in enumerate(doc):
        stored = (
            page_texts[page_idx]
            if page_texts and page_idx < len(page_texts)
            else None
        )
        if not stored:
            continue

        # Quick check: skip pages where no change phrase is present
        if not any(orig in stored for orig, _ in changes):
            continue

        blocks = _ocr_page_blocks(page)
        if not blocks:
            continue

        # ── Split stored text into paragraphs ──
        # First split on blank lines (double newlines)
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", stored) if p.strip()]
        # If the text uses single newlines only, fall back to line-based split
        # but merge consecutive short lines into paragraphs (min 80 chars)
        if len(paragraphs) <= 1:
            lines = [ln.strip() for ln in stored.split("\n") if ln.strip()]
            paragraphs = []
            current: list[str] = []
            for ln in lines:
                current.append(ln)
                if len(" ".join(current)) >= 80:
                    paragraphs.append(" ".join(current))
                    current = []
            if current:
                paragraphs.append(" ".join(current))

        # Separate leading labels (e.g. "NARRATIVE:") from content.
        # If a paragraph starts with a single-word label line followed
        # by the actual text, strip the label so we don't duplicate it
        # (the label is already visible in the original image).
        cleaned: list[str] = []
        for para in paragraphs:
            lines = para.split("\n", 1)
            if len(lines) == 2 and re.match(r"^[A-Z ]+:$", lines[0].strip()):
                cleaned.append(lines[1].strip())
            else:
                cleaned.append(para)
        paragraphs = cleaned

        # ── Identify affected paragraphs ──
        affected: list[tuple[str, str]] = []  # (original, debiased)
        for para in paragraphs:
            debiased = para
            changed = False
            for orig, repl in changes:
                if orig in debiased:
                    debiased = debiased.replace(orig, repl)
                    changed = True
            if changed:
                affected.append((para, debiased))

        if not affected:
            continue

        # ── Pre-compute normalised word sets for all blocks ──
        block_word_sets: list[set[str]] = []
        for block in blocks:
            block_word_sets.append(set(_normalize_ws(block["text"]).lower().split()))

        replaced_block_idx: set[int] = set()

        # ── For each affected paragraph, find matching OCR blocks ──
        for para, debiased_para in affected:
            para_words = set(_normalize_ws(para).lower().split())
            if not para_words:
                continue

            # Find blocks whose words overlap significantly with the paragraph.
            # Require BOTH a minimum percentage AND a minimum absolute count
            # to prevent small form-field blocks (1–4 words) from matching
            # a large narrative paragraph on common words like "S1", "Main", etc.
            matching_bis: list[int] = []
            for bi, bw_set in enumerate(block_word_sets):
                if bi in replaced_block_idx or not bw_set:
                    continue
                matching_count = len(bw_set & para_words)
                overlap = matching_count / len(bw_set)
                if overlap >= 0.4 and matching_count >= 5:
                    matching_bis.append(bi)

            # Fallback: find blocks containing a bias phrase via OCR text
            if not matching_bis:
                for bi, block in enumerate(blocks):
                    if bi in replaced_block_idx:
                        continue
                    bn = _normalize_ws(block["text"])
                    for orig, repl in changes:
                        if _normalize_ws(orig) in bn:
                            matching_bis.append(bi)
                            break

            if not matching_bis:
                logger.warning(
                    "Page %d: no OCR blocks matched paragraph: %.60s",
                    page.number, para,
                )
                continue

            # Merge matching blocks and render
            matched_blocks = [blocks[bi] for bi in matching_bis]
            merged = (
                _merge_blocks(matched_blocks) if len(matched_blocks) > 1 else matched_blocks[0]
            )
            _whiteout_and_render(page, merged, debiased_para)

            for bi in matching_bis:
                replaced_block_idx.add(bi)

            logger.info(
                "Page %d: replaced %d block(s) for paragraph (%.60s)",
                page.number, len(matching_bis), para,
            )


def _export_via_form_fields(
    doc: fitz.Document, changes: list[tuple[str, str]]
) -> bool:
    """Try to apply corrections via PDF form field widgets.

    Returns True if any form fields were found and updated.
    """
    updated = False
    for page in doc:
        for widget in page.widgets():
            value = widget.field_value
            if not value:
                continue
            modified = value
            changed = False
            for original, replacement in changes:
                if original in modified:
                    modified = modified.replace(original, replacement)
                    changed = True
            if changed:
                widget.field_value = modified
                widget.update()
                updated = True
                logger.debug("Updated field '%s' on page %d", widget.field_name, page.number)
    return updated


def _export_formatted_native(
    doc: fitz.Document,
    bias_changes: list[dict],
    entity_mapping: dict[str, str],
    page_texts: list[str] | None = None,
) -> None:
    """Apply bias corrections to a PDF by replacing the full narrative section."""
    changes = [
        (
            unmask_phrase(c["original_phrase"], entity_mapping),
            unmask_phrase(c["replacement_phrase"], entity_mapping),
        )
        for c in bias_changes
    ]

    if _export_via_form_fields(doc, changes):
        return

    logger.info("No form fields updated; using block-level replacement")
    _replace_affected_blocks(doc, changes, page_texts)


def _export_formatted_scanned(
    doc: fitz.Document,
    original_pdf_path: str | Path,
    bias_changes: list[dict],
    entity_mapping: dict[str, str],
    page_texts: list[str] | None = None,
) -> None:
    """Apply bias corrections on a scanned PDF using block-level replacement."""
    changes = [
        (
            unmask_phrase(c["original_phrase"], entity_mapping),
            unmask_phrase(c["replacement_phrase"], entity_mapping),
        )
        for c in bias_changes
    ]

    _replace_affected_blocks(doc, changes, page_texts)


def _add_searchable_text_layer(doc: fitz.Document) -> None:
    """Add an invisible OCR text layer to pages that lack extractable text.

    Renders each page to an image, runs Tesseract, and inserts every detected
    word as invisible text (render_mode=3).  The text is in the content stream
    so Adobe Acrobat (and other readers) can search / select it, but it doesn't
    change the visual appearance.

    Pages that already have substantial extractable text are skipped to avoid
    duplicate search hits.
    """
    import pytesseract
    from PIL import Image
    from pytesseract import Output

    dpi = 300
    scale = 72.0 / dpi

    for page in doc:
        # Skip pages that already have a usable text layer
        if len(page.get_text().strip()) > 50:
            continue

        pix = page.get_pixmap(dpi=dpi)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        ocr_data = pytesseract.image_to_data(img, output_type=Output.DICT)

        for i in range(len(ocr_data["text"])):
            word = ocr_data["text"][i].strip()
            if not word:
                continue

            x = ocr_data["left"][i] * scale
            y = ocr_data["top"][i] * scale
            h = ocr_data["height"][i] * scale
            fontsize = max(h * 0.85, 1)

            baseline_y = y + fontsize * 0.88
            page.insert_text(
                (x, baseline_y), word,
                fontname="helv", fontsize=fontsize,
                render_mode=3,  # invisible but searchable
            )

        logger.debug("Added searchable text layer to page %d", page.number)


def export_formatted_pdf(
    original_pdf_path: str | Path,
    bias_changes: list[dict],
    entity_mapping: dict[str, str],
    output_path: str | Path,
    is_scanned: bool = False,
    page_texts: list[str] | None = None,
) -> Path:
    """Export a formatted PDF with bias corrections applied in-place.

    For native PDFs with form fields: updates field values directly.
    For native PDFs without form fields: surgical per-phrase text replacement.
    For scanned PDFs: OCR bounding boxes to white-out and overlay text.

    After corrections, adds an invisible OCR text layer to any page that
    lacks extractable text, making the output searchable in Adobe Acrobat.
    """
    output_path = Path(output_path)
    doc = fitz.open(str(original_pdf_path))

    if not bias_changes:
        logger.info("No bias changes to apply; saving original PDF as-is.")
        doc.save(str(output_path))
        doc.close()
        return output_path

    if is_scanned:
        _export_formatted_scanned(
            doc, original_pdf_path, bias_changes, entity_mapping, page_texts=page_texts,
        )
    else:
        _export_formatted_native(
            doc, bias_changes, entity_mapping, page_texts=page_texts,
        )

    # Make non-searchable pages searchable via invisible OCR text layer
    _add_searchable_text_layer(doc)

    doc.save(str(output_path), garbage=3, deflate=True)
    doc.close()
    logger.info("Exported formatted PDF to %s", output_path)
    return output_path
